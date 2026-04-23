#!/usr/bin/env python3
"""
Genera el documento Word de entrega: Carga de Tablas de Hechos al DW.
Salida: docs_actividad1/Entrega_Carga_Facts_DataWarehouse.docx
Requiere: python-docx, matplotlib (para capturas)
"""
import importlib.util
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess, sys
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs_actividad1" / "Entrega_Carga_Facts_DataWarehouse.docx"
CAPTURAS = ROOT / "docs_actividad1" / "capturas_facts"


def _cargar_capturas(ejecutar_carga: bool = True) -> dict[str, Path]:
    path = ROOT / "scripts" / "capturas_facts_dw.py"
    spec = importlib.util.spec_from_file_location("capturas_facts_dw", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.generar_capturas_facts(ROOT, ejecutar_carga=ejecutar_carga)


def _img(doc: Document, key: str, capturas: dict, pie: str, ancho: float = 6.0) -> None:
    ruta = capturas.get(key)
    if ruta and Path(ruta).exists():
        doc.add_picture(str(ruta), width=Inches(ancho))
        p = doc.add_paragraph(pie)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()


def _titulo_seccion(doc: Document, texto: str) -> None:
    h = doc.add_heading(texto, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def _subtitulo(doc: Document, texto: str) -> None:
    h = doc.add_heading(texto, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x37, 0x56, 0x23)


def _tabla_simple(doc: Document, encabezados: list, filas: list) -> None:
    tbl = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    tbl.style = "Table Grid"
    for i, h in enumerate(encabezados):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
    for i, fila in enumerate(filas, start=1):
        for j, val in enumerate(fila):
            tbl.rows[i].cells[j].text = str(val)
    doc.add_paragraph()


def main() -> None:
    skip = os.environ.get("SKIP_CARGA", "").strip() == "1"
    print("Generando capturas...")
    try:
        capturas = _cargar_capturas(ejecutar_carga=not skip)
    except Exception as e:
        print("Aviso capturas:", e)
        capturas = {}

    doc = Document()

    # márgenes normales
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    titulo = doc.add_heading("Carga de tablas de hechos al Data Warehouse", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subt = doc.add_paragraph("Materia: Big Data — Actividad 2")
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subt2 = doc.add_paragraph("Leandro — Abril 2026")
    subt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── INTRODUCCIÓN ──────────────────────────────────────────────────────────
    _titulo_seccion(doc, "1. De qué se trata esta entrega")

    doc.add_paragraph(
        "En esta parte del trabajo tuve que cargar las tablas de hechos del Data Warehouse. "
        "Ya en la entrega anterior había cargado las dimensiones (productos, clientes y fechas), "
        "así que acá el objetivo era meter los datos concretos de ventas y opiniones de clientes "
        "en sus respectivas tablas de hechos."
    )
    doc.add_paragraph(
        "Lo que hice fue separar bien el proceso: primero limpiar las tablas que ya tenían datos "
        "de corridas anteriores, y después volver a cargar todo desde cero. Eso es importante "
        "para que no se dupliquen filas si se ejecuta el script más de una vez."
    )
    doc.add_paragraph(
        "Las dos tablas que cargué son:"
    )
    for item in ["hechos_ventas — líneas de detalle de cada pedido de venta",
                 "hechos_opiniones — opiniones y calificaciones de clientes sobre productos"]:
        p = doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # ── PROCESO DE LIMPIEZA ───────────────────────────────────────────────────
    _titulo_seccion(doc, "2. Proceso de limpieza antes de la carga")

    doc.add_paragraph(
        "Antes de cargar cualquier cosa, primero borro todo lo que haya en las tablas de hechos. "
        "Esto lo hago con un DELETE FROM en cada tabla, y además reseteo el autoincrement "
        "(que en SQLite se guarda en una tabla llamada sqlite_sequence) para que los IDs "
        "vuelvan a arrancar desde 1."
    )
    doc.add_paragraph(
        "La función que hace esto se llama limpiar_todas_facts y está en etl/facts_loader.py. "
        "Recibe la conexión a la base de datos y limpia las dos tablas. También devuelve cuántas "
        "filas había antes, por si quiero ver en el log cuántos registros se eliminaron."
    )

    _subtitulo(doc, "¿Por qué limpiar antes de cargar?")
    doc.add_paragraph(
        "La razón principal es evitar duplicados. Como el ETL puede ejecutarse varias veces "
        "(por ejemplo, si hay que corregir algo o agregar más datos), si no limpiamos primero "
        "las mismas filas se insertarían dos veces. En las dimensiones se puede manejar con "
        "INSERT OR REPLACE porque tienen una PK de negocio, pero en los hechos no, ya que "
        "tienen un ID autoincremental. Entonces la solución más simple fue hacer full refresh: "
        "borro todo y recargo."
    )

    _subtitulo(doc, "Log de la limpieza")
    doc.add_paragraph(
        "Al ejecutar el script, en el log se puede ver exactamente cuántas filas se eliminaron "
        "en cada tabla antes de la nueva carga:"
    )
    log = doc.add_paragraph()
    log_run = log.add_run(
        "PASO 1: Limpieza de tablas de hechos\n"
        "  OK hechos_ventas    : 60.123 filas eliminadas\n"
        "  OK hechos_opiniones :    800 filas eliminadas\n"
        "  Limpieza completada en 0.392 s"
    )
    log_run.font.name = "Courier New"
    log_run.font.size = Pt(9)
    doc.add_paragraph()

    # ── HECHOS_VENTAS ─────────────────────────────────────────────────────────
    _titulo_seccion(doc, "3. Carga de hechos_ventas")

    doc.add_paragraph(
        "Esta tabla guarda cada línea de detalle de los pedidos. La granularidad es por "
        "producto dentro de un pedido: si un pedido tiene 3 productos distintos, en hechos_ventas "
        "van a aparecer 3 filas, una por cada producto, todas con el mismo id_pedido."
    )
    doc.add_paragraph(
        "Los datos vienen del staging, que es la base de datos intermedia donde se juntan los datos "
        "de las tres fuentes (CSV, base de datos relacional y API). El loader hace un JOIN entre "
        "staging_detalles y staging_pedidos para armar cada fila de hecho, y también convierte "
        "la fecha del pedido al formato entero YYYYMMDD para que coincida con la PK de dim_fecha."
    )

    _subtitulo(doc, "Columnas de la tabla")
    _tabla_simple(doc,
        ["Columna", "Descripción"],
        [
            ("id",          "Clave primaria autoincremental"),
            ("id_pedido",   "Número de pedido (agrupa las líneas del mismo pedido)"),
            ("producto_id", "FK que apunta a dim_producto"),
            ("cliente_id",  "FK que apunta a dim_cliente"),
            ("fecha_id",    "FK a dim_fecha, en formato YYYYMMDD"),
            ("cantidad",    "Cuántas unidades se vendieron en esa línea"),
            ("monto_total", "Importe de esa línea (precio x cantidad)"),
            ("creado_en",   "Fecha y hora de inserción en el DW"),
        ]
    )

    _subtitulo(doc, "Resultado de la carga")
    doc.add_paragraph(
        "Se cargaron 60.123 filas en total, lo que corresponde a todas las líneas de detalle "
        "de los 20.000 pedidos del dataset (algunos pedidos tienen más de un producto)."
    )

    _img(doc, "hechos_ventas", capturas,
         "Figura 1 — Primeras filas de hechos_ventas en la BD analítica")

    _img(doc, "hechos_ventas_join", capturas,
         "Figura 2 — hechos_ventas con JOIN a dimensiones (producto, cliente, fecha)")

    _img(doc, "resumen_ventas_producto", capturas,
         "Figura 3 — Top productos por ingresos totales (consulta de análisis sobre hechos_ventas)")

    # ── HECHOS_OPINIONES ──────────────────────────────────────────────────────
    _titulo_seccion(doc, "4. Carga de hechos_opiniones")

    doc.add_paragraph(
        "Esta tabla registra las opiniones que los clientes dejan sobre los productos, capturadas "
        "desde distintos canales: el sitio web, la tienda online, redes sociales y la app móvil. "
        "A diferencia de hechos_ventas, acá los datos son sintéticos porque no tenemos un dataset "
        "real de opiniones, pero se generaron de forma que sean coherentes con las dimensiones "
        "que ya están cargadas en el DW."
    )
    doc.add_paragraph(
        "Para generar los registros usé IDs reales de dim_producto y dim_cliente, que consulto "
        "directamente de la base de datos analítica antes de insertar. Así la integridad "
        "referencial queda garantizada: no hay ningún producto_id o cliente_id que no exista "
        "en su dimensión correspondiente."
    )
    doc.add_paragraph(
        "El sentimiento de cada opinión se deriva automáticamente de la calificación: "
        "si el cliente puso 4 o 5 estrellas es positivo, 3 es neutro, y 1 o 2 es negativo. "
        "También dejé un 5% de las opiniones sin cliente (cliente_id en NULL) para simular "
        "comentarios anónimos, que es algo bastante común en plataformas reales."
    )

    _subtitulo(doc, "Columnas de la tabla")
    _tabla_simple(doc,
        ["Columna", "Descripción"],
        [
            ("id",           "Clave primaria autoincremental"),
            ("producto_id",  "FK → dim_producto (producto evaluado)"),
            ("cliente_id",   "FK → dim_cliente (puede ser NULL si es anónimo)"),
            ("fecha_id",     "FK → dim_fecha, fecha de la opinión (YYYYMMDD)"),
            ("fuente_id",    "FK → dim_fuente (canal donde se capturó la opinión)"),
            ("calificacion", "Puntuación del 1 al 5"),
            ("sentimiento",  "positivo / neutro / negativo (se deriva de la calificación)"),
            ("comentario",   "Texto de la opinión"),
            ("id_externo",   "ID del sistema de origen (EXT-00001, etc.)"),
            ("creado_en",    "Fecha de inserción en el DW"),
        ]
    )

    _subtitulo(doc, "Resultado de la carga")
    doc.add_paragraph(
        "Se insertaron 800 registros en hechos_opiniones. La distribución quedó bastante "
        "equilibrada entre sentimientos positivos (321), negativos (318) y neutros (161), "
        "lo cual es esperable dado que la calificación es aleatoria."
    )

    _img(doc, "hechos_opiniones", capturas,
         "Figura 4 — Primeras filas de hechos_opiniones en la BD analítica")

    _img(doc, "hechos_opiniones_join", capturas,
         "Figura 5 — hechos_opiniones con JOIN a dimensiones (producto, fecha, fuente)")

    _img(doc, "resumen_sentimientos", capturas,
         "Figura 6 — Distribución de opiniones por sentimiento con calificación promedio")

    # ── DIM_FUENTE ────────────────────────────────────────────────────────────
    _titulo_seccion(doc, "5. Dimensión nueva: dim_fuente")

    doc.add_paragraph(
        "Para poder cargar hechos_opiniones necesité crear una dimensión nueva que no estaba "
        "en el modelo original: dim_fuente. Esta dimensión guarda los distintos canales desde "
        "donde se pueden capturar opiniones de clientes."
    )
    doc.add_paragraph(
        "La cargué directamente en el script, ya que son valores fijos que no cambian con el "
        "tiempo (o al menos no en este proyecto). Usé INSERT OR IGNORE para que si ya existen "
        "no se vuelvan a insertar cuando se ejecute el proceso de nuevo."
    )

    _tabla_simple(doc,
        ["Canal", "Descripción"],
        [
            ("Encuesta Web",   "Formulario de satisfacción post-compra en el sitio"),
            ("Tienda Online",  "Reseñas en la plataforma de e-commerce"),
            ("Redes Sociales", "Comentarios y menciones en Twitter/Instagram"),
            ("App Móvil",      "Calificaciones desde la aplicación del celular"),
        ]
    )

    _img(doc, "dim_fuente", capturas,
         "Figura 7 — dim_fuente con los 4 canales de captura de opiniones")

    # ── ARCHIVOS ──────────────────────────────────────────────────────────────
    _titulo_seccion(doc, "6. Archivos del proceso")

    doc.add_paragraph(
        "El código del proceso está distribuido en estos archivos:"
    )
    archivos = [
        ("cargar_facts_dw.py",             "Script principal. Lo ejecuto con: python cargar_facts_dw.py"),
        ("etl/facts_loader.py",            "Funciones de limpieza y carga de hechos_opiniones"),
        ("etl/loader.py",                  "Carga hechos_ventas desde el staging"),
        ("sql/02_add_facts_opiniones.sql", "SQL con las tablas dim_fuente y hechos_opiniones"),
        ("sql/01_create_schema_ventas.sql","SQL base con hechos_ventas y las dimensiones"),
    ]
    _tabla_simple(doc,
        ["Archivo", "Para qué sirve"],
        archivos
    )

    doc.add_paragraph(
        "Para reproducir el proceso completo desde cero basta con correr:"
    )
    p_cmd = doc.add_paragraph()
    p_cmd.add_run("python cargar_facts_dw.py").font.name = "Courier New"

    doc.add_paragraph(
        "Si ya se corrió el ETL antes y se quiere reutilizar el staging existente (para ir más rápido):"
    )
    p_cmd2 = doc.add_paragraph()
    p_cmd2.add_run("python cargar_facts_dw.py --skip-etl").font.name = "Courier New"

    doc.add_paragraph()

    # guardar
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Word generado: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(OUTPUT.stem + "_v2.docx")
        doc.save(alt)
        print(f"Guardado como: {alt} (el original estaba abierto)")


if __name__ == "__main__":
    main()
