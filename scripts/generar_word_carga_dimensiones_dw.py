#!/usr/bin/env python3
"""
Genera el documento Word: carga de dimensiones al Data Warehouse (SQLite analítica).
Requiere: pip install python-docx matplotlib
Salida: docs_actividad1/Entrega_Carga_Dimensiones_DataWarehouse.docx

Las capturas de tablas se generan con scripts/capturas_carga_dw.py (PNG en
docs_actividad1/capturas_carga/). Por defecto se ejecuta run_etl antes.
Definir SKIP_ETL=1 en el entorno para reutilizar datos ya cargados (más rápido).
"""
import importlib.util
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess

    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs_actividad1" / "Entrega_Carga_Dimensiones_DataWarehouse.docx"
DOCS_ACT = ROOT / "docs_actividad1"
IMG_ARQ = DOCS_ACT / "ETL_Arquitectura.png"
IMG_FLUJO = DOCS_ACT / "ETL_Flujo.png"

REPO_HTTPS = "https://github.com/Lean-AZ/Proyecto_Ventas_ETL2"


def _cargar_modulo_capturas():
    path = ROOT / "scripts" / "capturas_carga_dw.py"
    spec = importlib.util.spec_from_file_location("capturas_carga_dw", path)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def _add_captura_tabla(doc: Document, ruta: Path | None, pie: str, ancho_pulg: float = 6.2) -> None:
    if ruta is not None and ruta.exists():
        doc.add_paragraph(pie)
        doc.add_picture(str(ruta), width=Inches(ancho_pulg))
        doc.add_paragraph()


def main() -> None:
    try:
        capt_mod = _cargar_modulo_capturas()
        run_etl = os.environ.get("SKIP_ETL", "").strip() != "1"
        capturas: dict[str, Path] = capt_mod.generar_capturas(ROOT, ejecutar_etl=run_etl)
    except Exception as e:
        print("Aviso: no se pudieron generar capturas automáticas:", e)
        capturas = {}

    doc = Document()

    title = doc.add_heading(
        "Carga de dimensiones al Data Warehouse", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subt = doc.add_paragraph(
        "Sistema de análisis de ventas — informe de la parte de carga al modelo analítico"
    )
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Portada", level=1)
    doc.add_paragraph(
        "Entrega: descripción del proceso de carga de dimensiones (y hechos) hacia la base analítica."
    )
    doc.add_paragraph(
        "El proyecto es un ETL en Python con SQLite; no montamos un servidor aparte "
        "porque la idea era poder correr todo en la PC y que el profe pueda clonar el repo y probarlo."
    )
    doc.add_paragraph("Fecha: abril de 2026.")
    doc.add_paragraph()
    doc.add_paragraph(
        "El código y los scripts están en GitHub por si quieren revisar línea por línea: "
        + REPO_HTTPS
    )
    doc.add_paragraph(
        "El PDF de entrega lo armé exportando este Word (Guardar como PDF); el contenido es el mismo."
    )
    doc.add_paragraph()

    doc.add_heading("1. Qué es lo que estoy documentando", level=1)
    doc.add_paragraph(
        "En la materia nos pidieron dejar escrito cómo queda la carga hacia el Data Warehouse. "
        "En la práctica, mi “warehouse” es un SQLite que guardamos como `data/ventas_analitica.db`; "
        "la ruta sale del `config/config.json`. Ahí van las tablas tipo estrella: dimensiones de "
        "fecha, producto y cliente, y después los hechos de ventas."
    )
    doc.add_paragraph(
        "La parte de código que hace la magia después del staging está en `etl/loader.py`. "
        "`run_etl.py` primero mete todo en staging y recién al final llama a `load_from_staging`, "
        "así que si algo falla en la carga, al menos ya tengo los datos intermedios para debuguear."
    )
    doc.add_paragraph()

    doc.add_heading("2. Cómo encaja la carga (staging → analítica)", level=1)
    doc.add_paragraph(
        "Cuando terminan los extractores, lo que queda “limpio” para el modelo está en "
        "`data/staging.db`: productos, clientes, pedidos y detalles, cada uno en su tabla "
        "(`staging_productos`, etc.). Desde ahí el loader abre la analítica, corre el SQL del "
        "esquema por si cambió algo, y empieza a pasar filas."
    )
    doc.add_paragraph(
        "No es nada exótico: es leer de un SQLite y escribir en otro. Lo que sí tuve cuidado "
        "es en no mezclar responsabilidades: el loader no vuelve a tocar CSV ni API, solo staging."
    )
    if IMG_ARQ.exists():
        doc.add_paragraph(
            "Abajo dejo el diagrama general del pipeline (extracción + destino) por si ayuda a ubicarse:"
        )
        doc.add_picture(str(IMG_ARQ), width=Inches(5.5))
        doc.add_paragraph()
    if IMG_FLUJO.exists():
        doc.add_paragraph("Y el flujo ordenado del ETL (mismo repo, exportado como imagen):")
        doc.add_picture(str(IMG_FLUJO), width=Inches(5.2))
        doc.add_paragraph()

    doc.add_heading("2.1 Evidencia: staging (antes del pasaje al modelo)", level=2)
    doc.add_paragraph(
        "Estas capturas las generé automáticamente leyendo el SQLite de staging "
        "después de correr el ETL (son tablas reales, no un mock de diseño)."
    )
    _add_captura_tabla(
        doc,
        capturas.get("staging_productos"),
        "Figura — Muestra de filas en staging_productos (origen de dim_producto):",
    )

    doc.add_heading("3. Flujo de la carga (resumido)", level=1)
    doc.add_paragraph(
        "El orden más o menos es este. Algunos pasos los saqué tal cual del código para no inventar:"
    )
    flujo = [
        "Leo rutas de staging y analítica desde el config.",
        "Abro la analítica y ejecuto `sql/01_create_schema_ventas.sql` (tablas dim_*, hechos, índices, FKs).",
        "Armo `dim_fecha` con un bucle de fechas (INSERT OR IGNORE) para no duplicar días si ya existían.",
        "Pasaje directo de `staging_productos` → `dim_producto` con INSERT OR REPLACE por id de producto.",
        "Igual con clientes: `staging_clientes` → `dim_cliente`.",
        "En hechos hice full refresh: borro todo `hechos_ventas` y lo vuelvo a llenar con un JOIN entre detalle y pedido.",
        "La fecha del pedido la paso a entero YYYYMMDD y esa es la `fecha_id` que apunta a `dim_fecha`.",
        "Commit, cierro conexiones, y el `run_etl.py` ya había dejado logueado el tiempo que tardó.",
    ]
    for item in flujo:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph()

    doc.add_heading("4. Tablas que terminan cargándose", level=1)
    doc.add_paragraph(
        "Abajo detallo qué guarda cada una y para qué la usamos en el modelo. "
        "Incluyo hechos porque sin eso las dimensiones quedan “colgadas”."
    )

    doc.add_heading("4.1 dim_fecha", level=2)
    doc.add_paragraph(
        "Es la dimensión de calendario. Acá algo importante: no la llenamos desde staging, "
        "sino con una función `_rellenar_dim_fecha` en `loader.py` que genera día por día "
        "(en el código dejé un rango tipo 2020–2026, se puede tocar si hace falta más rango)."
    )
    doc.add_paragraph(
        "El id del día lo armé como número entero YYYYMMDD para que sea fácil unir con los hechos "
        "sin andar parseando strings en cada consulta. También guardé año, mes, trimestre, semana ISO, "
        "y los nombres del mes/día en español porque para reportes suele venir bien."
    )
    dim_fecha_cols = [
        ("id", "PK numérica; formato YYYYMMDD (un día = un id)"),
        ("fecha", "Texto en formato ISO (YYYY-MM-DD), única"),
        ("anio, mes, trimestre", "Enteros para agrupar en reportes"),
        ("semana", "Semana ISO (la que usa isocalendar en Python)"),
        ("dia_semana", "1 = lunes … 7 = domingo"),
        ("nombre_mes, nombre_dia", "Texto en español, más legible en tablas"),
    ]
    tbl = doc.add_table(rows=1 + len(dim_fecha_cols), cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Columna"
    tbl.rows[0].cells[1].text = "Descripción"
    for i, (c, d) in enumerate(dim_fecha_cols, start=1):
        tbl.rows[i].cells[0].text = c
        tbl.rows[i].cells[1].text = d
    doc.add_paragraph()
    _add_captura_tabla(
        doc,
        capturas.get("dim_fecha"),
        "Figura — Muestra de dim_fecha en ventas_analitica.db (calendario cargado):",
    )

    doc.add_heading("4.2 dim_producto", level=2)
    doc.add_paragraph(
        "Viene de lo que quedó en `staging_productos` (eso a su vez mezcla lo que entró por CSV, "
        "API o la BD chica, según cómo corrió el merge en staging). Uso INSERT OR REPLACE con el "
        "mismo id de negocio: si el producto cambió de precio o categoría, pisa la fila vieja."
    )
    doc.add_paragraph(
        "En el SQL del esquema también están `creado_en` / `actualizado_en`; no las toqué mucho "
        "en el loader, pero quedan ahí por si después queremos auditar."
    )
    rows_p = [
        ("id", "PK; es el mismo id de producto que venimos usando en staging"),
        ("nombre", "Nombre del producto (obligatorio en el esquema)"),
        ("categoria", "Rubro / tipo, puede venir vacío"),
        ("precio", "Número decimal"),
        ("stock", "Cantidad en inventario (entero)"),
    ]
    tbl2 = doc.add_table(rows=1 + len(rows_p), cols=2)
    tbl2.style = "Table Grid"
    tbl2.rows[0].cells[0].text = "Columna"
    tbl2.rows[0].cells[1].text = "Descripción"
    for i, (c, d) in enumerate(rows_p, start=1):
        tbl2.rows[i].cells[0].text = c
        tbl2.rows[i].cells[1].text = d
    doc.add_paragraph()
    _add_captura_tabla(
        doc,
        capturas.get("dim_producto"),
        "Figura — Muestra de dim_producto tras la carga desde staging:",
    )

    doc.add_heading("4.3 dim_cliente", level=2)
    doc.add_paragraph(
        "Misma idea que producto: leo `staging_clientes` y hago REPLACE por id. Sirve para cortar "
        "las ventas por ciudad/país o ver datos de contacto si hace falta. No es un CRM, pero alcanza "
        "para el modelo de práctica."
    )
    rows_c = [
        ("id", "PK; id de cliente igual al de staging"),
        ("nombre_completo", "Nombre tal como llegó normalizado"),
        ("email, telefono", "Datos de contacto (a veces vienen incompletos)"),
        ("ciudad, pais", "Para mapas o rankings por país sin complicarse"),
    ]
    tbl3 = doc.add_table(rows=1 + len(rows_c), cols=2)
    tbl3.style = "Table Grid"
    tbl3.rows[0].cells[0].text = "Columna"
    tbl3.rows[0].cells[1].text = "Descripción"
    for i, (c, d) in enumerate(rows_c, start=1):
        tbl3.rows[i].cells[0].text = c
        tbl3.rows[i].cells[1].text = d
    doc.add_paragraph()
    _add_captura_tabla(
        doc,
        capturas.get("dim_cliente"),
        "Figura — Muestra de dim_cliente tras la carga desde staging:",
    )

    doc.add_heading("4.4 hechos_ventas (tabla de hechos)", level=2)
    doc.add_paragraph(
        "Acá guardamos cada línea de detalle de pedido (producto + cantidad + monto). "
        "Elegí recargarla entera en cada corrida del ETL: hago DELETE y después inserto de nuevo "
        "todo lo que salió del JOIN entre `staging_detalles` y `staging_pedidos`. Es simple; "
        "en un proyecto grande capaz haría incremental, pero para el tamaño del dataset no valía la pena."
    )
    doc.add_paragraph(
        "Las FK van a producto, cliente y fecha. La `fecha_id` la calculo sacando los guiones "
        "de la fecha del pedido para que calce con el id de `dim_fecha`."
    )
    rows_h = [
        ("id", "Surrogate autoincremental (cada fila de hecho nueva)"),
        ("id_pedido", "Sirve para agrupar líneas del mismo pedido"),
        ("producto_id", "Apunta a dim_producto"),
        ("cliente_id", "Apunta a dim_cliente"),
        ("fecha_id", "Apunta a dim_fecha (YYYYMMDD)"),
        ("cantidad", "Cuántas unidades en esa línea"),
        ("monto_total", "Plata de esa línea"),
    ]
    tbl4 = doc.add_table(rows=1 + len(rows_h), cols=2)
    tbl4.style = "Table Grid"
    tbl4.rows[0].cells[0].text = "Columna"
    tbl4.rows[0].cells[1].text = "Descripción"
    for i, (c, d) in enumerate(rows_h, start=1):
        tbl4.rows[i].cells[0].text = c
        tbl4.rows[i].cells[1].text = d
    doc.add_paragraph()
    _add_captura_tabla(
        doc,
        capturas.get("hechos_ventas"),
        "Figura — Muestra de hechos_ventas (líneas de detalle enlazadas a dimensiones):",
    )

    doc.add_heading("5. Dónde verlo en el repo", level=1)
    doc.add_paragraph(
        "Si algo no cierra con lo que escribí arriba, lo más directo es mirar estos archivos:"
    )
    for arch, desc in [
        ("sql/01_create_schema_ventas.sql", "Definición de tablas, FKs e índices."),
        ("etl/loader.py", "Toda la lógica de carga + relleno de fechas."),
        ("run_etl.py", "Acá se llama al loader al final del pipeline."),
        ("etl/staging.py", "Cómo quedan armadas las tablas staging_* antes del pasaje."),
    ]:
        doc.add_paragraph(f"• {arch} — {desc}")
    doc.add_paragraph()

    doc.add_heading("6. GitHub", level=1)
    doc.add_paragraph(
        "Link del repositorio para revisión (clonar, correr `run_etl.py`, etc.):"
    )
    doc.add_paragraph(REPO_HTTPS)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        dest = OUTPUT
    except PermissionError:
        alt = OUTPUT.with_name(OUTPUT.stem + "_actualizado.docx")
        doc.save(alt)
        dest = alt
        print(
            "Aviso: no se pudo sobrescribir el .docx (¿abierto en Word?). "
            f"Se guardó una copia: {alt}"
        )
    print(f"Documento Word generado: {dest}")


if __name__ == "__main__":
    main()
