#!/usr/bin/env python3
"""
Genera el documento Word de la Actividad 1 (ETL) a partir del contenido definido.
Requiere: pip install python-docx
"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).resolve().parent.parent / "docs_actividad1" / "Entrega_Actividad1_ETL.docx"
DOCS_ACT = Path(__file__).resolve().parent.parent / "docs_actividad1"
IMG_ARQ = DOCS_ACT / "ETL_Arquitectura.png"
IMG_FLUJO = DOCS_ACT / "ETL_Flujo.png"


def main():
    doc = Document()
    doc.add_heading("Actividad 1 – Desarrollo del Proceso ETL (Extracción)", 0)
    doc.add_paragraph("Documento técnico – Sistema de Análisis de Ventas")
    doc.add_paragraph()
    doc.add_paragraph("Repositorio: https://github.com/Lean-AZ/Proyecto_Ventas_ETL2")
    doc.add_paragraph()

    doc.add_heading("1. Diagrama de arquitectura", level=1)
    doc.add_paragraph(
        "La arquitectura incluye: Servicio ETL (Python, run_etl.py), fuentes (CSV, BD relacional, API REST), "
        "staging (data/staging.db), BD analítica (data/ventas_analitica.db) y dashboard como componente futuro."
    )
    if IMG_ARQ.exists():
        doc.add_picture(str(IMG_ARQ), width=Inches(5.5))
    doc.add_paragraph()
    doc.add_heading("Atributos de calidad", level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Atributo"
    h[1].text = "Cómo se garantiza"
    data = [
        ("Rendimiento", "Procesamiento por lotes; logging con time.perf_counter()."),
        ("Escalabilidad", "Nuevas fuentes = nuevo Extractor; config en config.json."),
        ("Seguridad", "config.json no versionado con secretos; solo config.json.example en repo."),
        ("Mantenibilidad", "Capas: extractors, staging, loader, config, logging; interfaz Extractor."),
    ]
    for i, (a, b) in enumerate(data, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    doc.add_paragraph()

    doc.add_heading("2. Diagrama de flujo del proceso ETL", level=1)
    doc.add_paragraph(
        "Flujo: Cargar config y logger → Extracción CSV → Extracción BD → Extracción API REST → "
        "Escribir en staging → DataLoader lee staging → Transformar (mapeo, fecha_id) → Cargar en BD analítica → Logs → Fin."
    )
    if IMG_FLUJO.exists():
        doc.add_picture(str(IMG_FLUJO), width=Inches(4))
    doc.add_paragraph()

    doc.add_heading("3. Justificación de decisiones técnicas", level=1)
    for p in [
        "Python: misma separación de responsabilidades e interfaz Extractor que en .NET.",
        "SQLite: portabilidad; rutas en config.json; en producción podría usarse PostgreSQL/SQL Server.",
        "Interfaz Extractor: cada fuente implementa extract(); se añaden fuentes sin tocar el orquestador.",
        "Config en JSON: rutas y URL centralizadas; credenciales no versionadas.",
        "Logging: módulo logging de Python (equivalente ILogger); consola y opcionalmente logs/etl.log.",
    ]:
        doc.add_paragraph(p, style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("4. Evidencia del código", level=1)
    evidence = [
        ("etl/extractors/base.py", "Clase abstracta Extractor con extract() -> dict."),
        ("etl/extractors/csv_extractor.py", "CsvExtractor: csv.DictReader, products, customers, orders, order_details."),
        ("etl/extractors/db_extractor.py", "DatabaseExtractor: SQLite fuente_ventas.db, pedidos y detalle_pedidos."),
        ("etl/extractors/api_extractor.py", "ApiExtractor: GET a API REST productos/clientes."),
        ("etl/staging.py", "Tablas staging_* y write_to_staging / merge_into_staging."),
        ("etl/loader.py", "Lectura staging, dim_fecha, carga en dim_* y hechos_ventas."),
        ("run_etl.py", "Orquestación: config, logger, tres extractores, staging, loader, time.perf_counter()."),
        ("etl/logger.py", "setup_logger() con formato y handlers consola/archivo."),
    ]
    for arch, desc in evidence:
        doc.add_paragraph(f"{arch}: {desc}", style="List Bullet")
    doc.add_paragraph("Código completo en el repositorio GitHub.")
    doc.add_paragraph()

    doc.add_heading("5. Repositorio", level=1)
    doc.add_paragraph("https://github.com/Lean-AZ/Proyecto_Ventas_ETL2")
    doc.save(OUTPUT)
    print(f"Documento Word generado: {OUTPUT}")


if __name__ == "__main__":
    main()
